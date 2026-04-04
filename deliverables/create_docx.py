#!/usr/bin/env python3
"""
Fiserv Digital Pay: Full Strategy Analysis - Word Document Generator
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

NAVY = RGBColor(0x00, 0x3B, 0x5C)
ORANGE = RGBColor(0xFF, 0x66, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = DARK

# Configure heading styles
for i, (size, color) in enumerate([(24, NAVY), (18, NAVY), (14, NAVY), (12, ORANGE)], 1):
    h_style = doc.styles[f'Heading {i}']
    h_style.font.name = 'Calibri'
    h_style.font.size = Pt(size)
    h_style.font.color.rgb = color
    h_style.font.bold = True

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003B5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_para(text, style='Normal', bold=False, size=None, color=None, space_after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
    return p

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('FISERV DIGITAL PAY')
run.font.size = Pt(36)
run.font.color.rgb = NAVY
run.font.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Crypto & Stablecoin Acceptance Strategy')
run.font.size = Pt(22)
run.font.color.rgb = ORANGE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Full Analysis Document')
run.font.size = Pt(16)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CommerceHub  |  Finxact  |  FIUSD  |  INDX')
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('May 2026  |  CONFIDENTIAL')
run.font.size = Pt(12)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prepared for Fiserv Executive Leadership & Investor Day')
run.font.size = Pt(11)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    'Part 1: Wallet Provider Deep Analysis',
    '    1A. Coinbase -- Full Ecosystem Analysis',
    '    1B. Mesh -- Full Platform Analysis',
    '    1C. Flexa -- Full Network Analysis',
    '    1D. Heron -- Full Assessment',
    '    1E. Comprehensive Comparative Analysis',
    '    1F. Wallet Recommendation',
    'Part 2: MPP Build-vs-Implement Analysis',
    'Part 3: Fiserv Strategic Positioning',
    'Part 4: Disruptive Prototypeable Ideas',
    'Part 5: E2E Prototype Architecture',
    'Part 6: Risk Assessment',
    'Appendix A: Executive Presentation Playbook',
    'Appendix B: Answered Questions & Decisions',
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if not item.startswith('    '):
        run.bold = True
        run.font.color.rgb = NAVY

doc.add_page_break()

# ============================================================
# CONTEXT
# ============================================================
doc.add_heading('Context', level=1)
add_para('Fiserv Commerce Hub is developing a crypto/stablecoin acceptance strategy across its merchant ecosystem. The company has a unique vertically integrated stack -- Finxact (banking core), FIUSD (stablecoin on Solana), INDX (real-time USD settlement), and CommerceHub/Clover (merchant acceptance) -- that no competitor possesses.')
add_para('Critical discovery: Fiserv is already a founding member of the x402 Foundation (alongside Coinbase, Stripe, Visa, Mastercard, Google, AWS, Adyen, Shopify) as of April 2, 2026.', bold=True, color=ORANGE)

doc.add_page_break()

# ============================================================
# PART 1: WALLET PROVIDERS
# ============================================================
doc.add_heading('Part 1: Wallet Provider Deep Analysis', level=1)

# 1A COINBASE
doc.add_heading('1A. Coinbase -- Full Ecosystem Analysis', level=2)

doc.add_heading('Product Portfolio (April 2026)', level=3)
add_table(
    ['Product', 'Purpose', 'Maturity', 'Relevance'],
    [
        ['Coinbase Business (fka Commerce)', 'Merchant crypto acceptance', 'Transitioning', 'HIGH'],
        ['Wallet-as-a-Service (WaaS)', 'White-label embedded wallets', 'Production', 'CRITICAL'],
        ['Coinbase Prime', 'Institutional custody + trading', 'Mature', 'HIGH'],
        ['CDP (Developer Platform)', 'x402, AgentKit, Smart Wallet', 'Production', 'CRITICAL'],
        ['Base L2', 'Low-cost settlement layer', 'Production', 'HIGH'],
    ]
)
doc.add_paragraph()

doc.add_heading('Wallet Connectivity', level=3)
add_para('Coinbase is NOT a walled garden for payment acceptance. Via WalletConnect integration, 700+ external wallets (MetaMask, Trust Wallet, Ledger, Rainbow, Phantom, etc.) can pay at a Coinbase-powered merchant checkout. $400B moved over WalletConnect in 2025; 72% in stablecoins. However, the settlement and fiat off-ramp side is Coinbase-only -- that is the strategic lock-in point.')

doc.add_heading('WaaS Architecture', level=3)
add_bullet('2-of-2 MPC threshold ECDSA -- key shares split between user device and Coinbase server')
add_bullet('MPC library open-sourced March 2025')
add_bullet('Full white-label: partners control all branding and UX')
add_bullet('Supports 8 chains: Base, Ethereum, Optimism, Arbitrum, Polygon, Avalanche, BNB, Zora')
add_bullet('Pricing: $0.005 per wallet operation (creation, signing, broadcast), first 500/month free')
add_bullet('Custom Stablecoins (Dec 2025): partners can issue branded stablecoins backed 1:1 by USDC with revenue share')

doc.add_heading('Coinbase Prime (Custody)', level=3)
add_bullet('$370B+ in assets under custody -- largest regulated crypto custodian globally')
add_bullet('Qualified Custodian under SEC rules')
add_bullet('SOC 1 Type 2 AND SOC 2 Type 2 audited by Deloitte')
add_bullet('Insurance: $320M via Lloyd\'s of London syndicate')
add_bullet('98% cold storage architecture')
add_bullet('Key client: BlackRock (iShares Bitcoin Trust custody)')

doc.add_heading('Base L2 Performance', level=3)
add_bullet('Transaction cost: ~$0.002 per USDC transfer')
add_bullet('Flashblocks (July 2025): 200-millisecond block times')
add_bullet('DeFi TVL: $4.63B (46% of all L2 market)')
add_bullet('Revenue: $75.4M in 2025 (62% of all L2 revenue)')
add_bullet('Native USDC (Circle-issued, not bridged)')

doc.add_heading('x402 Protocol Role', level=3)
add_bullet('Coinbase CREATED x402, donated to Linux Foundation April 2, 2026')
add_bullet('119M cumulative transactions on Base, 35M on Solana, ~$600M annualized volume')
add_bullet('Zero protocol fees. Free tier: 1,000 txns/month, then $0.001/txn')
add_bullet('AgentKit: Plug-and-play AI agent wallet operations with TEE-secured keys')
add_bullet('Paymaster: 15M+ gasless transactions across 50+ apps')

doc.add_heading('Security & Compliance', level=3)
add_table(
    ['Dimension', 'Detail'],
    [
        ['Licenses', '49 US states (MSB/MTL), EU MiCA, Singapore MAS, Japan FSA'],
        ['Federal charter', 'OCC National Trust Charter: conditional approval April 2, 2026'],
        ['Certifications', 'SOC 1 Type 2, SOC 2 Type 2 (Deloitte)'],
        ['Insurance', '$320M Lloyd\'s syndicate'],
        ['Cold storage', '98% of assets'],
        ['Major incident', 'May 2025: insider breach, 69,461 accounts. No crypto lost. $180-400M cost'],
        ['Regulatory', 'SEC lawsuit dismissed (Feb 2025). NYDFS $50M AML fine settled'],
    ]
)
doc.add_paragraph()

doc.add_heading('Pros', level=3)
pros = [
    'Only provider offering white-label MPC wallets for Finxact ledger',
    'Created x402 -- deepest protocol expertise',
    '$370B custody + $320M insurance -- institutional-grade security',
    'Custom Stablecoins service could power FIUSD issuance',
    '700+ wallets can pay via WalletConnect',
    'Base L2 offers $0.002 transaction costs with 200ms block times',
    'AgentKit + Paymaster solve the agentic commerce plumbing',
    'SOC 1 + SOC 2 by Deloitte -- audit trail for enterprise procurement',
]
for p in pros:
    add_bullet(p)

doc.add_heading('Cons', level=3)
cons = [
    'Direct competitive conflict: Coinbase Business competes with merchant payment providers',
    'Vendor lock-in risk: MPC key shares, Base sequencer, settlement',
    'No in-store/POS solution: QR codes only, no NFC or terminal integration',
    'Geographic gaps: Coinbase Business is US/Singapore only',
    'Uptime concerns: Near-daily minor outages; major outages during peaks',
    'Insider breach precedent: May 2025 incident',
    'Missing: No ISO 8583, no multi-acquirer routing, no recurring billing',
    'Settlement lock-in: fiat off-ramp is Coinbase-only',
]
for c in cons:
    add_bullet(c)

doc.add_heading('Fiserv Fit Assessment', level=3)
add_bullet('E-commerce: 9/10', bold_prefix='')
add_bullet('In-store: 3/10', bold_prefix='')
add_bullet('Finxact integration: 9/10', bold_prefix='')
add_bullet('Strategic risk: 6/10', bold_prefix='')

doc.add_page_break()

# 1B MESH
doc.add_heading('1B. Mesh -- Full Platform Analysis', level=2)

doc.add_heading('Company Profile', level=3)
add_table(
    ['Dimension', 'Detail'],
    [
        ['Founded', '2020 (as Front Financial, Inc.)'],
        ['Rebrand', 'Mid-2023 to Mesh Connect, Inc.'],
        ['Founders', 'Bam Azizi (CEO), Adam Israel (COO)'],
        ['Total funding', '$200M+ across 5 rounds'],
        ['Latest round', 'Series C, Jan 2026, $75M led by Dragonfly Capital'],
        ['Valuation', '$1B (unicorn)'],
        ['Key investors', 'Paradigm, Dragonfly, PayPal Ventures, Coinbase Ventures, Consensys'],
        ['Employees', 'Est. 100-250'],
    ]
)
doc.add_paragraph()

doc.add_heading('Core Technology: Wallet Aggregation', level=3)
add_bullet('300+ platforms (exchanges, self-custody wallets, brokerages)')
add_bullet('Single unified API -- integrate once, access all 300+ platforms')
add_bullet('Non-custodial: never holds funds or private keys')
add_bullet('SmartFunding (patent-pending): Combines up to 5 funding sources into a single transaction')
add_bullet('100+ tokens supported across Ethereum, Solana, Bitcoin, Polygon')

doc.add_heading('Payment Product: Mesh Pay', level=3)
add_bullet('Shopify: Native app on Shopify App Store (March 2025)')
add_bullet('Apple Pay integration for crypto payments')
add_bullet('Shift4 Partnership: 200,000+ merchants across 45+ countries')
add_bullet('AI Wallet (MAI Agent): Natural language purchasing, supports Google AP2')
add_bullet('Pricing: Not publicly disclosed -- custom enterprise only')
add_bullet('In-store/POS: NONE. Zero proprietary in-store capability')

doc.add_heading('Security & Compliance', level=3)
add_table(
    ['Dimension', 'Status'],
    [
        ['SOC 2 Type II', 'Certified'],
        ['Custody model', 'Non-custodial'],
        ['KYC/AML', 'Delegated to underlying exchanges'],
        ['Money transmitter license', 'NOT CONFIRMED'],
        ['Security incidents', 'None known'],
    ]
)
doc.add_paragraph()

doc.add_heading('Pros', level=3)
for p in ['Broadest wallet/exchange coverage (300+ platforms)', 'SmartFunding is genuinely unique (patent-pending)', 'Shift4 distribution: 200K+ merchants in 45+ countries', 'Apple Pay bridge for crypto', 'SOC 2 Type II certified', 'AP2 support for agentic commerce']:
    add_bullet(p)

doc.add_heading('Cons', level=3)
for c in ['Zero in-store/POS capability', 'Aggregation fragility: 300+ API dependencies', 'No chargeback/dispute/refund mechanism', 'No card fallback', 'Opaque pricing', 'No confirmed money transmitter license']:
    add_bullet(c)

doc.add_heading('Fiserv Fit', level=3)
add_bullet('E-commerce: 8/10. In-store: 2/10. Finxact integration: 5/10. Strategic risk: 5/10')

doc.add_page_break()

# 1C FLEXA
doc.add_heading('1C. Flexa -- Full Network Analysis', level=2)
add_para('Note: Flexa has been REMOVED from the recommended architecture due to critical risks identified below.', bold=True, color=ORANGE)

doc.add_heading('Company Profile', level=3)
add_table(
    ['Dimension', 'Detail'],
    [
        ['Founded', '2018, New York City'],
        ['Total funding', '$14.1M (Series A -- 7+ years, still Series A)'],
        ['Employees', '~43-45'],
        ['Revenue model', '~1% fee, ~0.90% redistributed to AMP stakers'],
    ]
)
doc.add_paragraph()

doc.add_heading('AMP Collateral-Backed Payment Guarantee', level=3)
add_para('Flexa uses a genuinely novel architecture where AMP tokens are locked as collateral to guarantee merchant payment. This enables zero chargebacks and instant fiat settlement. However, AMP lost 98% of its value from 2021-2024.')

doc.add_heading('Why Flexa Was Removed', level=3)
for c in [
    'AMP token lost 98% of value -- collateral model depends on a collapsing token',
    'Only $14.1M raised after 7+ years (still Series A)',
    'US/Canada only -- unusable for global CommerceHub merchants',
    'No confirmed Clover/NCR/Verifone integration',
    'Revenue model unclear (fees go to AMP stakers, not the company)',
    'Better alternative: Build in-store natively on Clover using Finxact FIUSD wallets + INDX',
]:
    add_bullet(c)

doc.add_page_break()

# 1D HERON
doc.add_heading('1D. Heron -- Full Assessment', level=2)
add_para('Note: Heron has been REMOVED from the evaluation. The entity cannot be verified as a going concern.', bold=True, color=ORANGE)

doc.add_heading('Critical Finding: Entity Confusion', level=3)
add_para('There are THREE distinct companies named Heron: tryheron.com (stablecoin orchestrator), heronfinance.com (private credit), and herondata.io (AI document automation). Only tryheron.com is relevant.')

doc.add_heading('tryheron.com Assessment', level=3)
add_table(
    ['Dimension', 'Status'],
    [
        ['Founded', 'Unknown'],
        ['Funding', 'Unknown -- not on Crunchbase, PitchBook, or Tracxn'],
        ['Team', 'Unknown -- no team page'],
        ['Named customers', 'Zero'],
        ['Press coverage', 'Zero third-party mentions'],
        ['API documentation', 'None public'],
        ['Security certifications', 'None disclosed'],
    ]
)
doc.add_paragraph()
add_para('Recommendation: For B2B stablecoin treasury orchestration in Phase 2, evaluate Bridge (Stripe, $1.1B acquisition) or BVNK (Mastercard, ~$1.8B acquisition), or build natively on Finxact + INDX.')

doc.add_page_break()

# 1E COMPARATIVE
doc.add_heading('1E. Comprehensive Comparative Analysis', level=2)

doc.add_heading('Scoring Matrix: 7 Requirements', level=3)
add_table(
    ['#', 'Requirement', 'Coinbase', 'Mesh', 'Flexa*', 'Heron*'],
    [
        ['1', 'Crypto pay-in (global)', '9/10', '9/10', '7/10', '1/10'],
        ['2', 'Stablecoin pay-in (global)', '9/10', '8/10', '7/10', '3/10'],
        ['3', 'E-commerce + in-store', '5/10', '5/10', '7/10', '1/10'],
        ['4', 'Stablecoin disbursement', '8/10', '4/10', '2/10', '5/10'],
        ['5', 'Wallet storage on Finxact', '9/10', '4/10', '2/10', '2/10'],
        ['6', 'Stablecoin-to-fiat', '9/10', '8/10', '8/10', '4/10'],
        ['7', 'Any crypto-to-fiat (global)', '7/10', '9/10', '8/10', '1/10'],
        ['TOTAL', '', '56/70', '47/70', '41/70', '17/70'],
    ],
    col_widths=[0.3, 2.5, 1.0, 1.0, 1.0, 1.0]
)
doc.add_paragraph()
add_para('*Flexa and Heron have been removed from the recommended architecture.', size=9, color=GRAY)

doc.add_heading('Security Comparison', level=3)
add_table(
    ['Dimension', 'Coinbase', 'Mesh', 'Flexa', 'Heron'],
    [
        ['SOC 2 Type II', 'Yes (Deloitte)', 'Yes', 'No', 'No'],
        ['Insurance', '$320M Lloyd\'s', 'None', 'AMP collateral', 'None'],
        ['Assets custody', '$370B+', '$0 (non-custodial)', 'N/A', '$0'],
        ['US Licenses', '49 states + OCC', 'None confirmed', 'FinCEN MSB + MTLs', 'None'],
        ['Int\'l Licenses', 'EU MiCA, SG, JP', 'None', 'None', 'None'],
        ['KYC/AML', 'Full internal', 'Delegated', 'Full internal', 'Unknown'],
    ]
)
doc.add_paragraph()

doc.add_page_break()

# 1F RECOMMENDATION
doc.add_heading('1F. Wallet Recommendation', level=2)

add_para('No single provider covers all 7 requirements. The gaps are structural, not incremental. A layered architecture behind a Fiserv-owned abstraction layer is required.')

doc.add_heading('Recommended Architecture', level=3)
add_para('Fiserv Digital Pay (Merchant-Facing Brand)\n    -> CommerceHub Crypto Orchestration Layer\n        -> Coinbase WaaS+CDP (Primary: wallet, x402, custody)\n        -> Mesh Pay (Reach: 300+ wallets, SmartFunding)\n        -> Finxact + FIUSD + INDX (Owned: settlement, in-store via Clover)', size=10)

doc.add_heading('Provider Role Assignment', level=3)
add_table(
    ['Provider', 'Role', 'Why', 'Contract'],
    [
        ['Coinbase WaaS + CDP', 'Primary wallet, x402, custody, conversion', 'Only white-label MPC wallets + x402 creator + SOC 2 + $370B custody', '90-day termination. <40% volume cap'],
        ['Mesh Pay', 'Multi-wallet reach, SmartFunding, e-commerce', '300+ integrations + SmartFunding + Apple Pay', 'Secondary acceptance connector'],
        ['Finxact + FIUSD + INDX', 'Owned wallet, stablecoin, settlement, in-store', 'FDIC-insured. Clover in-store path. Full fallback', 'Internal -- proprietary stack'],
    ]
)
doc.add_paragraph()

doc.add_heading('Key Design Principles', level=3)
for p in [
    'Abstraction layer is non-negotiable: merchants see "Fiserv Digital Pay" only',
    'No single provider >40% of volume',
    '90-day termination clauses on all vendor contracts',
    'Card-rail fallback on every transaction',
    'Finxact-native for in-store (Clover + FIUSD + INDX)',
    'Circuit breaker on stablecoin depeg (>0.5% deviation)',
]:
    add_bullet(p)

doc.add_page_break()

# ============================================================
# PART 2: MPP
# ============================================================
doc.add_heading('Part 2: MPP Build-vs-Implement Analysis', level=1)

add_para('BOTTOM LINE: DO NOT BUILD A PROPRIETARY PROTOCOL.', bold=True, color=ORANGE, size=14)
add_para('Every unique Fiserv capability can be delivered as implementation-layer value on top of open standards. Building a proprietary protocol would fragment the ecosystem, cost 12-18 months, and solve no problem that implementing existing protocols cannot.')

doc.add_heading('Why Engineering Wants to Build (And Why They Should Not)', level=2)
for p in [
    'Fiserv already committed publicly to x402 Foundation (April 2, 2026)',
    'Adyen (closest competitive analog) is NOT building their own protocol',
    'Protocol adoption requires ecosystem (22+ members vs. zero)',
    'No identified capability requires a proprietary protocol',
]:
    add_bullet(p)

doc.add_heading('x402 vs MPP: Complementary, Not Competitive', level=2)
add_table(
    ['Dimension', 'x402', 'MPP'],
    [
        ['State model', 'Stateless (per-request)', 'Stateful (session/escrow)'],
        ['Settlement', 'On-chain per transaction', 'Batched via escrow'],
        ['Fiat support', 'Crypto/stablecoins only', 'Hybrid: fiat + crypto + Lightning'],
        ['Best for', 'One-off purchases, API access', 'Streaming, subscriptions, sessions'],
        ['Micropayments', 'Gas cost per txn adds up', 'Efficient batching'],
        ['Governance', 'Linux Foundation (22+ members)', 'IETF draft + Stripe/Tempo'],
    ]
)
doc.add_paragraph()

doc.add_heading('What Engineering Should Build: "CommerceHub Agent Pay"', level=2)
add_para('A protocol-agnostic agentic commerce gateway that:')
for p in [
    'Translates between protocols: x402, MPP, AP2, TAP',
    'Defaults to FIUSD on Solana for lowest-cost settlement',
    'Bank-verified agent identity via Finxact KYC layer',
    'Real-time USD settlement via INDX regardless of protocol',
    '6M merchant auto-enablement via software update',
    'Fiserv Agent SDK (TypeScript + Python) abstracts protocol selection',
]:
    add_bullet(p)

add_para('This is the exact same strategic pattern that made CommerceHub successful for card payments: don\'t own the network, own the orchestration.', bold=True)

doc.add_page_break()

# ============================================================
# PART 3: STRATEGIC POSITIONING
# ============================================================
doc.add_heading('Part 3: Fiserv Strategic Positioning', level=1)

doc.add_heading('The Unique Vertical Integration', level=2)
add_para('Fiserv has something NO other payment processor has: a closed-loop stablecoin settlement stack from bank to merchant.')
add_para('Bank issues FIUSD (Finxact) -> Consumer holds FIUSD -> Consumer pays merchant (CommerceHub/Clover) -> Real-time USD settlement (INDX) -> Merchant receives USD INSTANTLY', size=10, bold=True)

doc.add_heading('Competitive Moat Analysis', level=3)
add_table(
    ['Competitor', 'Banking Core', 'Own Stablecoin', 'Settlement', 'Merchants', 'Full Stack?'],
    [
        ['Fiserv', 'Finxact', 'FIUSD', 'INDX (1,100+ banks)', '6M (CommerceHub+Clover)', 'YES'],
        ['Stripe', 'No', 'No (Bridge/USDC)', 'Via Bridge', 'Stripe Checkout', 'No'],
        ['Adyen', 'No', 'No', 'No', 'Adyen platform', 'No'],
        ['Worldpay', 'No', 'No', 'Via BVNK', 'Worldpay platform', 'No'],
        ['PayPal', 'No', 'PYUSD', 'No', 'PayPal/Venmo', 'No'],
    ]
)
doc.add_paragraph()

doc.add_heading('Blue Ocean Positioning', level=2)
add_table(
    ['Action', 'Factor'],
    [
        ['ELIMINATE', 'Crypto volatility risk (auto-convert to USD via INDX)'],
        ['REDUCE', 'Integration complexity (one API toggle in CommerceHub)'],
        ['RAISE', 'Regulatory confidence (FDIC-insured, bank-embedded stablecoin)'],
        ['CREATE', 'Bank-merchant stablecoin loop (banks earn more, merchants pay less)'],
    ]
)
doc.add_paragraph()

doc.add_heading('Bank Economics', level=2)
add_bullet('Current bank interchange revenue: ~50bps per card transaction')
add_bullet('Proposed FIUSD transaction revenue: 60-80bps per FIUSD transaction')
add_bullet('Merchant savings: pays 80-100bps (FIUSD) vs 200-300bps (card interchange)')
add_para('Result: Bank earns MORE, merchant pays LESS, because the card network\'s share is eliminated.', bold=True, color=ORANGE)
add_para('FIUSD balances must be structured as bank deposits with a stablecoin wrapper (counts toward bank balance sheet) to address deposit cannibalization concerns.')

doc.add_page_break()

# ============================================================
# PART 4: PROTOTYPEABLE IDEAS
# ============================================================
doc.add_heading('Part 4: Disruptive Prototypeable Ideas', level=1)

ideas = [
    ('Priority 1: "Merchant Yield Sweep"',
     'SMB merchants on Clover have idle cash in settlement accounts earning nothing.',
     'AI treasury agent auto-sweeps idle merchant balances into yield-bearing FIUSD positions on Finxact.',
     'CommerceHub has the transaction data to predict cash needs. Finxact holds both deposits and yield positions. INDX provides instant liquidity.',
     'Turns Fiserv from cost center (processing fees) to value center (merchant wealth builder).'),
    ('Priority 2: "Pay-by-Agent" x402 Commerce',
     'AI shopping agents need to make payments programmatically. Card rails require human auth (3DS, CVV).',
     'CommerceHub exposes x402-compatible Agent Checkout endpoint. Agent discovers, negotiates, pays in FIUSD, receives receipt.',
     'x402 founding member. CommerceHub is checkout orchestration. FIUSD is settlement. Finxact manages spending limits.',
     'First working demo of AI agent purchasing on production payment platform. Distribution: 6M merchants with one update.'),
    ('Priority 3: "Instant Supplier Pay"',
     'Restaurant orders supplies by card (2-3% fee). Distributor waits 2-5 days for settlement.',
     'AI agent monitors inventory, auto-generates POs, pays suppliers instantly in FIUSD. Supplier gets USD in seconds.',
     'Clover sees merchant sales AND inventory. Finxact holds both accounts. CommerceHub orchestrates.',
     'Restaurant saves $1,200/month via early-payment discounts. Distributor improves cash flow.'),
    ('Priority 4: "Cross-Border Instant Settlement"',
     'Cross-border card processing costs 3-5%. Supplier payments cost 2-4% in FX, take 3-5 days.',
     'CommerceHub detects cross-border, converts to FIUSD at real-time FX rate, settles via INDX instantly.',
     'INDX connects 1,100+ banks. Finxact holds multi-currency positions. CommerceHub handles routing.',
     '$190T annual market. 90% cost reduction. Settlement in seconds, not days.'),
    ('Priority 5: "Stablecoin Loyalty Loop"',
     'Loyalty programs are fragmented, expensive (1-3% of revenue), with poor redemption.',
     'Merchants issue branded stablecoin-denominated loyalty credits on Finxact multi-asset ledger.',
     'Finxact natively represents merchant loyalty tokens alongside FIUSD. Clover handles POS issuance.',
     'Loyalty tokens are ledgered assets on a regulated banking core, not toy blockchain tokens.'),
]

for title, problem, solution, why_fiserv, impact in ideas:
    doc.add_heading(title, level=2)
    add_para('Problem: ' + problem, bold=True)
    add_para('Solution: ' + solution)
    add_para('Why only Fiserv: ' + why_fiserv)
    add_para('Impact: ' + impact, color=ORANGE)

doc.add_page_break()

# ============================================================
# PART 5: PROTOTYPE ARCHITECTURES
# ============================================================
doc.add_heading('Part 5: E2E Prototype Architecture', level=1)
add_para('All 4 prototypes built to full scope in parallel. FIUSD is live in production. Finxact engineering is aligned. 10-13 dedicated engineers across 4 teams. Each prototype is a standalone demo application.', bold=True)

# ---- PROTOTYPE 1 ----
doc.add_heading('Prototype 1: Merchant Yield Sweep', level=2)
add_para('Timeline: 6 weeks  |  Team: 2-3 engineers  |  Standalone demo app', bold=True, color=ORANGE)

doc.add_heading('Problem Statement', level=3)
add_para('SMB merchants on Clover have idle cash sitting in settlement accounts earning absolutely nothing. The average Clover merchant holds $15,000-$50,000 in settlement balances at any given time. Enterprise merchants on CommerceHub have dedicated treasury teams managing this manually at significant cost -- often $200K-$500K/year in treasury management overhead. Meanwhile, stablecoin yield opportunities (3-5% APY on FIUSD via Finxact) go untapped because no automated bridge exists between merchant settlement accounts and yield-bearing stablecoin positions.')
add_para('The result: billions of dollars in aggregate merchant settlement balances sit dormant across Fiserv\'s 6 million merchant locations, earning nothing for the merchants and generating no incremental revenue for Fiserv.')

doc.add_heading('Solution', level=3)
add_para('An AI-powered treasury agent that automatically sweeps idle merchant settlement balances into yield-bearing FIUSD positions on Finxact. The agent uses CommerceHub transaction data to build a multi-factor cash flow prediction model, ensuring the merchant always has sufficient liquidity for upcoming obligations (payroll, rent, supplier payments, tax remittances) while putting excess funds to work.')
add_para('Key capabilities:')
for item in [
    'Multi-factor cash flow prediction using 6+ months of CommerceHub settlement history',
    'ML-based seasonal pattern detection (holiday spikes, slow periods, day-of-week patterns)',
    'Obligation scheduling: integrates known fixed costs (payroll dates, rent due dates) into liquidity forecasting',
    'Automatic sweep: when predicted balance exceeds 3-day projected outflows, excess moves to FIUSD yield position on Finxact',
    'Automatic unsweep: when upcoming obligations require more liquidity, FIUSD converts back to USD instantly via INDX',
    'Clover dashboard widget showing real-time earnings, sweep history, and projected monthly yield',
    'Configurable risk tolerance: merchants can set conservative (5-day buffer), moderate (3-day), or aggressive (1-day) sweep thresholds',
]:
    add_bullet(item)

doc.add_heading('How We Implement It', level=3)
add_para('Technical Architecture:')
add_para('Clover POS / CommerceHub (Settlement Data Feed)\n  -> AI Treasury Agent (Cash Flow Predictor)\n    -> 3-day rolling forecast with seasonal adjustment\n    -> Obligation scheduler (payroll, rent, taxes)\n    -> Sweep decision engine\n  -> Finxact Ledger\n    -> Demand Deposit Account <-> FIUSD Yield Position Account\n    -> Atomic sweep/unsweep via Finxact position transfer APIs\n  -> INDX Settlement (Real-time USD liquidity when needed)', size=10)

add_para('Implementation stack:')
for item in [
    'Backend: Python/FastAPI service consuming CommerceHub settlement webhooks',
    'ML model: scikit-learn gradient boosting for cash flow prediction, trained on simulated 6-month merchant transaction history',
    'Finxact integration: REST API calls to create FIUSD yield position accounts, execute position transfers (sweep/unsweep)',
    'INDX simulation: Mock service replicating INDX real-time settlement behavior (actual INDX integration in production)',
    'Frontend: React dashboard embedded in Clover merchant portal -- earnings widget, sweep history chart, threshold configuration',
    'Database: PostgreSQL for merchant transaction history, sweep logs, and prediction model state',
    'Scheduling: Cron-based sweep evaluation every 15 minutes during business hours',
]:
    add_bullet(item)

doc.add_heading('Build Plan', level=3)
add_table(['Week', 'Deliverable'], [
    ['1-2', 'Cash flow prediction model using CommerceHub settlement history. Train on simulated 6-month merchant data. Build multi-factor predictor with seasonal adjustment. Set up PostgreSQL schema for transaction history and sweep logs.'],
    ['3', 'Finxact integration: Create FIUSD yield position account type via Finxact APIs. Build sweep/unsweep API calls with atomic position transfer. Implement configurable threshold engine (conservative/moderate/aggressive).'],
    ['4', 'INDX simulation layer: Build mock INDX service for instant USD liquidity. Implement automatic unsweep when obligations approach. Build reconciliation service tracking yield accrual.'],
    ['5', 'Clover dashboard widget: React component showing "Your idle cash earned $X this month." Sweep history chart with daily granularity. Threshold configuration UI. Projected annual yield calculator.'],
    ['6', 'Full demo scenario: Restaurant merchant "Mario\'s Pizzeria" earns $847 in one month on idle settlement balances. Walk through daily sweep decisions, show payroll unsweep, display cumulative earnings.'],
], col_widths=[0.6, 5.7])
doc.add_paragraph()

doc.add_heading('Demo Script for Investor Day', level=3)
add_para('"Today, a Clover merchant\'s settlement sits idle earning nothing. We just turned their dormant capital into a revenue stream. This restaurant -- Mario\'s Pizzeria -- earned $847 last month without doing anything. One toggle in their Clover dashboard. Our AI treasury agent analyzes their cash flow, predicts when they need money for payroll and suppliers, and sweeps the excess into yield-bearing FIUSD on Finxact. When they need cash for Friday payroll, it\'s back instantly via INDX. We\'re not just processing their payments anymore -- we\'re growing their money. Across 6 million merchants, this is a multi-billion dollar yield opportunity."')

doc.add_heading('Full Scope (No Cuts)', level=3)
add_para('Multi-factor cash flow prediction with ML-based seasonal patterns, full Finxact API integration for position management, INDX settlement simulation, interactive Clover dashboard with earnings visualization, configurable sweep thresholds, and obligation scheduling.')

doc.add_page_break()

# ---- PROTOTYPE 2 ----
doc.add_heading('Prototype 2: Pay-by-Agent x402 Commerce', level=2)
add_para('Timeline: 10 weeks  |  Team: 3-4 engineers  |  Standalone demo app', bold=True, color=ORANGE)

doc.add_heading('Problem Statement', level=3)
add_para('AI shopping agents (Claude, GPT, custom enterprise agents) are increasingly being used to browse, compare, and purchase products on behalf of consumers. However, current card payment rails require human-in-the-loop authentication at the point of payment -- 3D Secure challenges, CVV entry, biometric confirmation. This fundamentally breaks the autonomous agentic flow. An AI agent that can find the perfect product in 2 seconds but then must wait 30 seconds for a human to authenticate the payment defeats the purpose of agentic commerce.')
add_para('Additionally, card rails impose 2-3% interchange fees on every transaction, create T+1 to T+3 settlement delays for merchants, and provide no native mechanism for machine-to-machine payment negotiation. The emerging x402 protocol (HTTP 402 Payment Required) solves all of these problems, but no major payment orchestrator has implemented it at scale.')

doc.add_heading('Solution', level=3)
add_para('CommerceHub exposes an x402-compatible "Agent Checkout" endpoint that enables AI agents to discover payment requirements, negotiate terms, and complete payments programmatically -- all within a single HTTP request cycle, with no human authentication needed for pre-authorized spending amounts.')
add_para('Key capabilities:')
for item in [
    'x402 v2 full specification implementation on CommerceHub -- HTTP 402 response with payment instructions (price, accepted tokens, recipient address, chain)',
    'Multi-chain support: FIUSD + USDC on both Solana and Base (Coinbase L2)',
    'EIP-3009 (authorization transfer) and Permit2 signature verification for stablecoin payments',
    'Bank-verified agent identity via Finxact KYC tier -- agents authenticated through regulated banks get higher spending limits than wallet-only agents',
    'Configurable spending guardrails: per-transaction limits, daily limits, merchant category restrictions, all managed through Finxact',
    'Fiserv Agent SDK (TypeScript + Python) providing fetchWithPayment() wrapper that abstracts the entire x402 flow',
    'Real-time USD settlement to merchant via INDX -- merchant receives USD, never touches stablecoin',
    'CommerceHub dashboard integration: x402 agent transactions appear alongside Visa/MC transactions with "Agent" label',
    'Cryptographic receipt generation for agent audit trails',
]:
    add_bullet(item)

doc.add_heading('How We Implement It', level=3)
add_para('Technical Architecture:')
add_para('AI Shopping Agent (Claude/GPT/Custom)\n  Uses Fiserv Agent SDK (fetchWithPayment())\n  -> HTTP GET /product (browse merchant catalog)\n  -> CommerceHub x402 Gateway\n    Returns: HTTP 402 Payment Required\n    + X-PAYMENT header: {price, FIUSD address, accepted_tokens, chain, expiry}\n  -> Agent signs payment (EIP-3009 signed payload)\n  -> Agent retries with X-PAYMENT header containing signed payload\n  -> CommerceHub Payment Verifier\n    - Validates EIP-3009/Permit2 signature\n    - Checks agent identity against Finxact KYC tier\n    - Verifies spending limits (per-txn, daily, category)\n    - Settles FIUSD/USDC on-chain (Solana or Base)\n  -> Dual output:\n    - Returns product/service + cryptographic receipt to agent\n    - INDX auto-converts stablecoin to USD, settles to merchant CommerceHub account', size=10)

add_para('Implementation stack:')
for item in [
    'x402 Gateway: Node.js/Express server implementing HTTP 402 spec v2 with payment instruction headers',
    'Payment Verifier: Solana web3.js + ethers.js for on-chain signature validation and settlement',
    'Finxact KYC Integration: REST API calls to verify agent identity tier and retrieve spending limits',
    'Fiserv Agent SDK: TypeScript package (@fiserv/agent-pay) and Python package (fiserv-agent-pay) wrapping fetchWithPayment()',
    'INDX Simulation: Mock service for stablecoin-to-USD conversion and merchant settlement',
    'Demo Agent: Claude-powered shopping agent using Anthropic API with tool_use for product discovery and x402 payment',
    'Merchant Dashboard: React app showing x402 transactions alongside simulated card transactions',
    'On-chain: Solana devnet for FIUSD, Base Sepolia testnet for USDC -- both with deployed test token contracts',
]:
    add_bullet(item)

doc.add_heading('Build Plan', level=3)
add_table(['Week', 'Deliverable'], [
    ['1-2', 'CommerceHub x402 endpoint: Implement HTTP 402 response with payment instructions per x402 spec v2. Set up Solana devnet and Base Sepolia testnet. Deploy test FIUSD and USDC token contracts. Build payment instruction generator.'],
    ['3-4', 'Payment verification service: Validate EIP-3009/Permit2 signatures. Build Finxact KYC integration for agent identity tier lookup. Implement spending limit checks (per-transaction, daily, merchant category). Build on-chain settlement execution.'],
    ['5-6', 'Fiserv Agent SDK (TypeScript + Python): Build fetchWithPayment() wrapper handling full x402 flow. Implement bank-verified identity attachment via Finxact KYC tier. Package as @fiserv/agent-pay (npm) and fiserv-agent-pay (pip).'],
    ['7-8', 'INDX integration: Build auto-conversion service (FIUSD/USDC -> USD). Implement real-time settlement to merchant CommerceHub account (simulated). Build reconciliation between on-chain settlement and merchant ledger.'],
    ['9', 'Demo agent: Build Claude-powered shopping agent using Anthropic API with tool_use. Agent browses product catalog, discovers x402 endpoint, pays with FIUSD, receives cryptographic receipt. Full end-to-end flow in under 3 seconds.'],
    ['10', 'CommerceHub dashboard: React app showing x402 transactions alongside card transactions. "Agent" label on agentic purchases. Transaction detail view with on-chain receipt, settlement status, and agent identity.'],
], col_widths=[0.6, 5.7])
doc.add_paragraph()

doc.add_heading('Demo Script for Investor Day', level=3)
add_para('"Watch this. I\'m going to ask an AI agent to buy me a pair of running shoes. [Live demo: Agent browses merchant catalog, finds Nike Air Max in size 11, discovers x402 payment endpoint, signs FIUSD payment on Solana, receives cryptographic receipt -- all in under 3 seconds, zero human authentication needed.] The merchant saw this sale in their CommerceHub dashboard just like any Visa transaction -- except settlement was instant, not T+2, and the fee was 0.1%, not 2.3%. Now imagine this at scale: 6 million Clover merchants, all agent-payable, enabled with a single software update. Stripe has to convince merchants one by one. We flip a switch."')

doc.add_heading('The Distribution Advantage', level=3)
add_para('Auto-enable x402+MPP endpoints on ALL 6 million Clover merchants and ALL CommerceHub enterprise merchants via a software update. Overnight, Fiserv becomes the largest x402 merchant network in the world. Every AI shopping agent that speaks x402 can immediately transact at every Fiserv merchant. This is the exact distribution advantage that made Visa ubiquitous -- except Fiserv can do it with a software push, not decades of physical card issuance.')

doc.add_heading('Full Scope (No Cuts)', level=3)
add_para('Multi-chain support (Solana + Base), x402 v2 full specification, Fiserv Agent SDK in both TypeScript and Python, Finxact KYC integration for bank-verified agent identity, configurable spending guardrails, INDX settlement simulation, CommerceHub dashboard with agent transaction visibility, and Claude-powered demo agent.')

doc.add_page_break()

# ---- PROTOTYPE 3 ----
doc.add_heading('Prototype 3: Instant Supplier Pay', level=2)
add_para('Timeline: 8 weeks  |  Team: 2-3 engineers  |  Standalone demo app', bold=True, color=ORANGE)

doc.add_heading('Problem Statement', level=3)
add_para('A typical restaurant on Clover places 15-30 supply orders per week. Today, each order follows this painful flow: the manager calls or emails the distributor, places the order manually, pays by credit card (2-3% interchange fee), and the distributor waits 2-5 business days for settlement. The restaurant also pays invoices on net-30 terms, missing early-payment discounts that distributors routinely offer (2% if paid within 10 days).')
add_para('The costs compound: 2-3% interchange on every order, $0 captured early-payment discounts, 4-8 hours/week of manual ordering and reconciliation labor, and periodic stockouts when reordering is delayed. For a restaurant doing $50K/month in supply purchases, this means ~$1,000-1,500/month in card fees alone, plus another $1,000+ in missed discounts and operational inefficiency.')

doc.add_heading('Solution', level=3)
add_para('An AI procurement agent embedded in Clover that monitors ingredient usage patterns, predicts reorder needs, auto-generates purchase orders, and pays suppliers instantly in FIUSD via CommerceHub. The supplier receives USD in their bank account via INDX in seconds, not days. Because the supplier gets paid instantly, the AI agent automatically negotiates and captures early-payment discounts.')
add_para('Key capabilities:')
for item in [
    'Clover POS integration: ingest real-time sales data and correlate with ingredient depletion rates',
    'Multi-ingredient prediction model: tracks top 20+ ingredients, learns usage patterns by day-of-week and seasonality',
    'Multi-supplier support: maintains supplier catalog with pricing, lead times, minimum order quantities, and discount terms',
    'Automated PO generation: creates purchase orders when predicted stock falls below reorder threshold',
    'Early-payment discount negotiation: AI agent automatically applies 2%/net-10 discount terms when paying within 24 hours vs net-30',
    'FIUSD B2B payment: instant payment from merchant Finxact account to supplier Finxact account via CommerceHub',
    'INDX settlement: supplier receives USD in their bank account in real-time (seconds, not days)',
    'Clover dashboard: procurement agent activity feed, cost savings tracker, supplier payment history',
    'Approval workflows: configurable auto-approval thresholds (e.g., auto-approve orders under $500, require manager approval above)',
]:
    add_bullet(item)

doc.add_heading('How We Implement It', level=3)
add_para('Technical Architecture:')
add_para('Clover POS (Tracks sales data + inventory levels)\n  -> AI Procurement Agent\n    - Monitors ingredient usage via sales data correlation\n    - Predicts reorder needs using gradient boosting model\n    - Generates POs with supplier catalog lookup\n    - Applies early-payment discount logic (2%/net-10 vs net-30)\n    - Routes approved POs to CommerceHub\n  -> CommerceHub (B2B Payment Rail)\n    - Executes FIUSD transfer from merchant Finxact account\n    - Routes to supplier Finxact account\n  -> Settlement:\n    - Merchant Finxact account debited (FIUSD)\n    - Supplier receives USD via INDX in seconds\n    - Automatic reconciliation in both merchant and supplier ledgers', size=10)

add_para('Implementation stack:')
for item in [
    'Backend: Python/FastAPI service with Clover REST API integration for sales data ingestion',
    'ML model: scikit-learn gradient boosting for ingredient depletion prediction, trained on simulated restaurant data',
    'Supplier service: PostgreSQL-backed catalog with pricing, lead times, MOQs, and payment terms',
    'PO engine: Rule-based PO generation with configurable reorder points and approval thresholds',
    'Finxact integration: REST API for FIUSD position transfers between merchant and supplier accounts',
    'INDX simulation: Mock settlement service for supplier fiat payouts',
    'Frontend: React dashboard for Clover merchant portal -- procurement feed, savings tracker, supplier management',
]:
    add_bullet(item)

doc.add_heading('Build Plan', level=3)
add_table(['Week', 'Deliverable'], [
    ['1-2', 'Clover inventory integration: Build sales data ingestion pipeline. Create ingredient depletion model for a restaurant use case (top 20 ingredients). Set up supplier catalog with 3 simulated distributors (produce, protein, dry goods).'],
    ['3-4', 'AI procurement agent: Auto-generate POs based on predicted depletion curves. Implement early-payment discount logic (2% discount if paid within 24h vs net-30). Build approval workflow engine with configurable thresholds.'],
    ['5-6', 'CommerceHub B2B payment: FIUSD transfer from merchant Finxact account to supplier Finxact account. INDX simulation settles supplier in USD in real-time. Build reconciliation service for both sides.'],
    ['7', 'Clover dashboard: Procurement agent activity feed showing auto-orders, payments, and savings. Cost savings tracker: "You saved $1,200 this month via early-payment discounts." Supplier payment history with settlement timestamps.'],
    ['8', 'Full demo scenario: Restaurant "Mario\'s Pizzeria" auto-orders chicken from distributor "Fresh Foods Inc." AI detects low stock, generates PO, pays in FIUSD, distributor gets USD in 3 seconds (not 5 days), restaurant earns 2% discount. Show month-end savings report.'],
], col_widths=[0.6, 5.7])
doc.add_paragraph()

doc.add_heading('Demo Script for Investor Day', level=3)
add_para('"This restaurant on Clover -- Mario\'s Pizzeria -- just saved $1,200 this month. How? Their AI procurement agent monitored ingredient usage, detected that chicken breast stock would run out by Thursday, auto-ordered from their distributor Fresh Foods Inc., and paid instantly in FIUSD via CommerceHub. The distributor got their money in 3 seconds instead of waiting 5 days. Because the distributor got paid instantly, they gave Mario\'s a 2% early-payment discount -- automatically. The restaurant saves money. The distributor improves cash flow. Both are on our rails. Nobody touched a card. And that 2-3% interchange fee? Gone."')

doc.add_heading('Full Scope (No Cuts)', level=3)
add_para('Full inventory prediction with ML model, multi-supplier support (3+ distributors), AI procurement agent with early-pay negotiation and approval workflows, B2B FIUSD payment flow via CommerceHub, INDX settlement simulation, and interactive Clover dashboard with savings tracking.')

doc.add_page_break()

# ---- PROTOTYPE 4 ----
doc.add_heading('Prototype 4: Cross-Border Instant Settlement', level=2)
add_para('Timeline: 12 weeks  |  Team: 2-3 engineers  |  Standalone demo app', bold=True, color=ORANGE)

doc.add_heading('Problem Statement', level=3)
add_para('Cross-border payments are a $190 trillion annual market riddled with inefficiency. When a buyer in Mexico purchases from a US merchant on CommerceHub, the traditional card flow imposes 3-5% in cross-border processing fees plus 2-4% in FX markup, and the merchant waits T+3 days for settlement through correspondent banking chains. For a $1,000 transaction, the merchant loses $47-90 in fees and waits 3 business days to receive funds.')
add_para('For outbound payments (a US merchant paying an international supplier), the costs are even worse: SWIFT transfers cost $25-50 per transaction, take 3-5 business days, and require manual reconciliation across banking intermediaries. The total friction cost of cross-border commerce is estimated at $120 billion annually globally.')

doc.add_heading('Solution', level=3)
add_para('CommerceHub automatically detects cross-border transactions (buyer currency differs from merchant currency) and offers a stablecoin settlement route as an alternative to traditional card rails. The buyer\'s local currency is converted to FIUSD at real-time FX rates, transmitted on Solana in seconds, and settled to the merchant\'s bank account in USD via INDX -- all in under 10 seconds, at approximately 0.5% total cost.')
add_para('Key capabilities:')
for item in [
    'Multi-corridor support: MXN (Mexico), EUR (Europe), GBP (UK) to USD conversions',
    'Real-time FX rate engine: pulls live rates from multiple liquidity providers, locks rate for 30-second validity window',
    'Automatic route comparison: for every cross-border transaction, calculates cost/speed on both card rails and stablecoin rails',
    'FIUSD bridge: converts local currency to FIUSD on Solana via liquidity provider APIs',
    'INDX real-time settlement: converts FIUSD to USD and settles to merchant\'s bank account in seconds',
    'Side-by-side comparison dashboard: shows merchant the actual cost/speed difference for each cross-border transaction',
    'Reconciliation engine: automated matching of FX conversion, on-chain transfer, and fiat settlement records',
    'Regulatory compliance: transaction monitoring, sanctions screening, and reporting for cross-border flows',
]:
    add_bullet(item)

doc.add_heading('How We Implement It', level=3)
add_para('Technical Architecture:')
add_para('International Buyer (pays in MXN/EUR/GBP)\n  -> CommerceHub Cross-Border Gateway\n    - Detects currency mismatch (buyer vs merchant)\n    - Calculates dual routing: card rails vs stablecoin rails\n    - Presents cost/speed comparison to merchant\n    - If stablecoin route selected:\n  -> FX Conversion Service\n    - Real-time rate from multiple liquidity providers\n    - Rate lock with 30-second validity window\n    - Converts local currency to FIUSD on Solana\n  -> INDX Settlement Network\n    - Receives FIUSD\n    - Converts to USD in real-time\n    - Settles to US merchant\'s bank account\n  -> Merchant receives USD instantly\n    - Cost: ~0.5% (vs 3-5% on card rails)\n    - Time: seconds (vs T+3 days on card rails)', size=10)

add_para('Implementation stack:')
for item in [
    'Cross-border detection: CommerceHub middleware analyzing buyer geolocation, IP, card BIN, and currency code',
    'FX engine: Python service aggregating rates from simulated liquidity providers (mocking real providers like CurrencyCloud, Wise, Circle)',
    'Rate lock service: Redis-backed rate cache with 30-second TTL and atomic conversion execution',
    'Solana integration: web3.js for FIUSD transfers on Solana devnet with SPL token program',
    'INDX simulation: Mock service for FIUSD-to-USD conversion and bank settlement',
    'Comparison dashboard: React app with side-by-side view for each cross-border transaction',
    'Reconciliation: PostgreSQL-backed ledger matching FX conversion, on-chain transfer, and fiat settlement records',
]:
    add_bullet(item)

doc.add_heading('Build Plan', level=3)
add_table(['Week', 'Deliverable'], [
    ['1-3', 'CommerceHub cross-border detection and routing engine: Identify transactions where buyer currency != merchant currency. Build dual-route calculator comparing card rails (3-5% fee, T+3) vs stablecoin rails (0.5%, seconds). Implement route selection logic.'],
    ['4-6', 'FX conversion service: Build rate aggregation from simulated liquidity providers for MXN, EUR, GBP corridors. Implement 30-second rate lock with Redis. Build FIUSD minting/conversion pipeline on Solana devnet.'],
    ['7-9', 'INDX integration: Build FIUSD->USD real-time settlement simulation. Implement reconciliation engine matching FX conversion to on-chain transfer to fiat settlement. Build transaction monitoring and compliance logging.'],
    ['10-11', 'Comparison dashboard: React app showing side-by-side view for each cross-border transaction. Card route: "$47.50 fee, 3 days, 4 intermediaries" vs Stablecoin route: "$5.00 fee, 3 seconds, direct." Monthly savings summary. Corridor analytics.'],
    ['12', 'Full demo: Live cross-border transaction from simulated Mexican buyer to US merchant. Show real-time FIUSD conversion, on-chain transfer, INDX settlement, and merchant receiving USD -- all in under 10 seconds. Display the $42.50 savings on a $1,000 transaction.'],
], col_widths=[0.6, 5.7])
doc.add_paragraph()

doc.add_heading('Demo Script for Investor Day', level=3)
add_para('"This US merchant on CommerceHub just received a payment from a buyer in Mexico. On traditional card rails, that\'s a 3-5% cross-border fee and 3-day settlement. Watch what happens on our stablecoin rails: [Live demo: CommerceHub detected the cross-border transaction, converted pesos to FIUSD at real-time rate, transferred on Solana in 400 milliseconds, and settled the merchant in USD through INDX.] Cost: 0.5%. Time: 3 seconds. The merchant saved $42.50 on a single $1,000 transaction. Now look at this dashboard -- it shows every cross-border transaction this month, with a running total of savings. Cross-border payments are a $190 trillion annual market. We just made it real-time and 90% cheaper."')

doc.add_heading('Full Scope (No Cuts)', level=3)
add_para('Multi-corridor support (MXN, EUR, GBP to USD), real-time FX conversion with rate locking, Solana devnet FIUSD transfers, INDX settlement simulation, side-by-side comparison dashboard with corridor analytics, reconciliation engine, and compliance logging.')

doc.add_page_break()

# ---- PROTOTYPE SUMMARY ----
doc.add_heading('Prototype Summary & Sequencing', level=2)

add_table(
    ['Prototype', 'Team', 'Weeks', 'Investor Day Narrative'],
    [
        ['1. Merchant Yield Sweep', '2-3 engineers', '6', '"We help merchants EARN money" -- new revenue model'],
        ['2. Pay-by-Agent x402', '3-4 engineers', '10', '"First AI agent purchasing on production rails" -- thought leadership'],
        ['3. Instant Supplier Pay', '2-3 engineers', '8', '"AI-driven B2B, zero card fees" -- SMB value story'],
        ['4. Cross-Border Settlement', '2-3 engineers', '12', '"$190T market, 90% cost reduction" -- massive TAM'],
    ],
    col_widths=[2.0, 1.2, 0.8, 3.0]
)
doc.add_paragraph()

add_para('Total engineering needed: 10-13 dedicated engineers across 4 parallel tracks. All prototypes are standalone demo applications with simulated Finxact and INDX integrations where sandbox access is not yet available. All demos run independently and can be presented sequentially during the Investor Day 10-minute slot.', bold=True)

doc.add_page_break()

# ============================================================
# PART 6: RISK ASSESSMENT
# ============================================================
doc.add_heading('Part 6: Risk Assessment', level=1)

add_table(
    ['Risk', 'Severity', 'Probability', 'Mitigation'],
    [
        ['Regulatory reversal', 'Catastrophic', 'Low-Med', 'FIUSD as deposit wrapper; remove wrapper if restricted'],
        ['Stripe protocol lock-in', 'High', 'Medium', '6M auto-enable + dual protocol (x402 + MPP)'],
        ['Bank client resistance', 'High', 'High', 'FIUSD = deposit + 60-80bps (> interchange); phased rollout'],
        ['Consumer adoption stall', 'Medium', 'High', 'Lead with B2B/agent use cases first'],
        ['Stablecoin depeg event', 'Catastrophic', 'Low', 'Circuit breaker at 0.5%; instant auto-conversion'],
        ['Internal BU misalignment', 'High', 'High', 'C-suite sponsor + cross-BU task force with P&L'],
        ['Coinbase competitive conflict', 'Medium', 'Medium', 'Abstraction layer + <40% volume cap + 90-day terms'],
    ],
    col_widths=[2.0, 1.2, 1.0, 3.0]
)
doc.add_paragraph()

add_para('The greatest risk is inaction. Stripe acquired Bridge for $1.1B. Worldpay partnered with BVNK. Adyen joined every major protocol foundation. The window for Fiserv to establish its unique position is narrowing.', bold=True)

doc.add_page_break()

# ============================================================
# APPENDIX A
# ============================================================
doc.add_heading('Appendix A: Executive Presentation Playbook', level=1)

add_para('10-minute Investor Day presentation flow:', bold=True)
steps = [
    '(1 min) The Opportunity: $2-4T stablecoin market. Sub-1% costs vs 2-3% interchange.',
    '(1 min) Our Unique Position: The vertically integrated stack diagram.',
    '(2 min) Live Demo 1 -- Yield Sweep: Clover dashboard, merchant earns yield.',
    '(2 min) Live Demo 2 -- Pay-by-Agent: AI agent purchases via x402.',
    '(2 min) Live Demo 3 -- Supplier Pay + Cross-Border combined demo.',
    '(1 min) Provider Architecture: Coinbase + Mesh behind Fiserv abstraction.',
    '(1 min) The Ask: Executive sponsor, cross-BU task force, 18-month roadmap.',
]
for s in steps:
    add_bullet(s)

doc.add_page_break()

# ============================================================
# APPENDIX B
# ============================================================
doc.add_heading('Appendix B: Answered Questions & Decisions', level=1)

add_table(
    ['#', 'Question', 'Answer', 'Action'],
    [
        ['1', 'Investor Day timing', '3rd week of May 2026', '7-week build with buffer'],
        ['2', 'Engineering allocation', '10-13 engineers confirmed', 'All 4 prototypes in parallel'],
        ['3', 'Finxact API access', 'Custom MCP connector available', 'Use for integration'],
        ['4', 'INDX sandbox', 'Simulate', 'Build simulation layer'],
        ['5', 'Coinbase CDP access', 'Sandbox available', 'Use for Pay-by-Agent prototype'],
        ['6', 'Deliverables', 'PPTX deck + Word document', 'Both generated'],
        ['7', 'Demo environment', 'Standalone demo apps', 'Build 4 independent applications'],
    ],
    col_widths=[0.3, 1.8, 2.0, 2.5]
)

# ============================================================
# SAVE
# ============================================================
# Add headers and footers
for section in doc.sections:
    header = section.header
    hp = header.paragraphs[0]
    hp.text = 'Fiserv Digital Pay | Crypto & Stablecoin Acceptance Strategy | CONFIDENTIAL'
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = 'May 2026 | Fiserv, Inc.'
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'

output_path = '/Users/ajnarasi/Documents/Work/Projects/stablecoin/deliverables/Fiserv_Crypto_Strategy_Full_Analysis.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
