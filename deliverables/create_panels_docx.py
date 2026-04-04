#!/usr/bin/env python3
"""
Fiserv Crypto Strategy: Business Panel + Spec Panel Evaluation Report
Export to Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

NAVY = RGBColor(0x00, 0x3B, 0x5C)
ORANGE = RGBColor(0xFF, 0x66, 0x00)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
TEAL = RGBColor(0x00, 0x80, 0x90)

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

for i, (size, color) in enumerate([(24, NAVY), (18, NAVY), (14, NAVY), (12, ORANGE)], 1):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.size = Pt(size)
    h.font.color.rgb = color
    h.font.bold = True


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003B5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            if i % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                cell._tc.get_or_add_tcPr().append(shading)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


def p(text, bold=False, size=None, color=None, space_after=6, align=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    run.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(space_after)
    if align: para.alignment = align
    return para


def bullet(text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = para.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
        r = para.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    else:
        r = para.add_run(text)
        r.font.name = 'Calibri'
        r.font.size = Pt(10)


def expert(name, framework, text):
    """Add an expert commentary block."""
    para = doc.add_paragraph()
    r = para.add_run(f'{name}')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    r.font.name = 'Calibri'
    r = para.add_run(f' ({framework}):')
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY
    r.font.name = 'Calibri'
    para.paragraph_format.space_after = Pt(2)
    p(text, size=10, space_after=10)


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p('FISERV DIGITAL PAY', bold=True, size=36, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Expert Panel Evaluations', size=22, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Business Strategy Panel + Technical Specification Panel', size=16, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

for _ in range(3):
    doc.add_paragraph()

p('Business Panel: Christensen, Porter, Drucker, Godin, Kim & Mauborgne, Collins, Taleb, Meadows, Doumont', size=11, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
p('Spec Panel: Fowler, Nygard, Newman, Wiegers, Adzic, Crispin, Hightower, Hohpe', size=11, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
p('May 2026  |  CONFIDENTIAL', size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    'Section 1: Business Panel Analysis',
    '    Idea 1: Merchant Yield Sweep',
    '    Idea 2: Pay-by-Agent x402 Commerce',
    '    Idea 3: Instant Supplier Pay',
    '    Idea 4: Cross-Border Instant Settlement',
    '    Overall Business Panel Ranking',
    '    Panel Consensus Recommendations',
    'Section 2: Technical Specification Panel Review',
    '    Prototype 1: Merchant Yield Sweep',
    '    Prototype 2: Pay-by-Agent x402 Commerce',
    '    Prototype 3: Instant Supplier Pay',
    '    Prototype 4: Cross-Border Instant Settlement',
    '    Cross-Cutting Technical Recommendations',
    '    Shared Infrastructure Design',
    '    Investor Day Technical Setup',
    'Section 3: Consolidated Action Items',
]
for item in toc:
    para = doc.add_paragraph()
    r = para.add_run(item)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    if not item.startswith('    '):
        r.bold = True
        r.font.color.rgb = NAVY

doc.add_page_break()

# ============================================================
# SECTION 1: BUSINESS PANEL
# ============================================================
doc.add_heading('Section 1: Business Panel Analysis', level=1)
p('Mode: DEBATE  |  All 9 experts convened', bold=True, color=ORANGE)
p('Experts: Clayton Christensen (Disruption), Michael Porter (Competitive Strategy), Peter Drucker (Management), Seth Godin (Marketing), W. Chan Kim & Renee Mauborgne (Blue Ocean), Jim Collins (Excellence), Nassim Taleb (Risk), Donella Meadows (Systems Thinking), Jean-luc Doumont (Communication)')

doc.add_page_break()

# ---- IDEA 1: YIELD SWEEP ----
doc.add_heading('Idea 1: Merchant Yield Sweep', level=2)

doc.add_heading('Expert Analysis', level=3)

expert('Clayton Christensen', 'Disruption Theory',
       'This is a textbook low-end disruption aimed at non-consumption. The job-to-be-done is clear: merchants are not hiring a "yield product" -- they are hiring something to make their idle money stop being idle. Today, 99% of SMB merchants consume ZERO treasury management services because existing solutions are priced for enterprises with $10M+ in liquid assets. The Clover merchant with $30K in settlement balances is a non-consumer of treasury services. This idea attacks non-consumption with a product that is simpler (one toggle), cheaper (zero setup), and more convenient (embedded in Clover). However, the yield SOURCE must be bulletproof -- if it comes from DeFi lending with smart contract risk, that is inappropriate for merchant operating capital at a regulated financial institution.')

expert('Michael Porter', 'Five Forces',
       'Supplier power: LOW (Fiserv controls entire value chain). Buyer power: LOW (merchants are fragmented, inertia of embedded defaults works for Fiserv). Substitutes: MEDIUM (bank sweep accounts exist but target enterprises). New entrants: LOW-MEDIUM (requires banking core + stablecoin + settlement + POS). Rivalry: LOW today. The competitive moat is genuinely strong -- no competitor can replicate the closed loop: Clover (data) + Finxact (ledger) + INDX (liquidity). Stripe cannot do this. Revenue potential: 6M merchants x $25K avg idle = $150B aggregate. At 4% yield with 50bps Fiserv take = $750M/year at full penetration. 5-10% adoption year 1-2 = $37-75M.')

expert('Peter Drucker', 'Management Philosophy',
       'The positioning must be relentlessly simple. Do not call this "yield sweep" or "treasury agent." Call it "Fiserv SmartCash" or "Money That Works" -- something a restaurant owner with 10 minutes of attention span understands. Deeper concern: this sits at the intersection of four BUs (Clover, CommerceHub, Finxact, digital assets). Who owns the P&L? Who is accountable when the AI sweeps too aggressively and a merchant can\'t make payroll? Appoint a SINGLE product owner with cross-BU authority.')

expert('Nassim Taleb', 'Risk Management',
       'The convexity profile is favorable but has hidden fragility. Upside: smooth and linear (more merchants, more yield). Downside: catastrophic and non-linear -- if ANY merchant has a liquidity shortfall because the AI swept too aggressively, the lawsuit and press coverage would damage Fiserv\'s banking relationships far more than the yield revenue is worth. NON-NEGOTIABLE safeguards: (1) Hard floor: never sweep below highest historical daily obligation + 20% buffer. (2) Real-time override: instant 100% unsweep, 24/7/365. (3) Insurance: Fiserv backstops any liquidity shortfall caused by the algorithm. (4) Gradual ramp: start with 5% of idle for first 30 days. WITHOUT these, I would vote to KILL this idea.')

expert('Seth Godin', 'Marketing',
       'The story is extraordinary: "We don\'t just process your payments. We grow your money." The dashboard widget showing "$847 earned this month" is the single most viral feature Fiserv could build -- every merchant will show it to other merchants. But rename it from "Merchant Yield Sweep" to something a restaurant owner would say out loud.')

expert('Kim & Mauborgne', 'Blue Ocean Strategy',
       'This creates genuine blue ocean. The strategy canvas for SMB treasury services has ZERO competitors addressing merchants with <$100K. ELIMINATE complexity (one toggle). REDUCE risk (hard floor + instant unsweep). RAISE convenience (embedded, zero effort). CREATE new value (merchants earning yield -- does not exist today).')

expert('Jim Collins', 'Good to Great',
       'Hedgehog Concept alignment is perfect. What can Fiserv be best at? Earning yield on merchant idle cash -- no one else has Clover data + Finxact ledger + INDX liquidity. The flywheel compounds: each merchant adds to the aggregate yield pool, enabling better rates, attracting more merchants.')

expert('Donella Meadows', 'Systems Thinking',
       'Highest-leverage intervention: INFORMATION FLOW. The dashboard makes the invisible visible -- merchants see what their idle cash is costing them (opportunity cost). The reinforcing feedback loop: merchants earn, tell other merchants, more adopt, aggregate pool grows, better rates, more earning. Virtuous cycle with network effects.')

doc.add_heading('Synthesis', level=3)
add_table(
    ['Dimension', 'Assessment'],
    [
        ['Strengths', 'Attacks non-consumption, embedded distribution (6M merchants), changes Fiserv from cost to value center, strong moat (closed-loop stack), viral dashboard, clear revenue model'],
        ['Weaknesses', 'Yield source must be transparent and insured, AI prediction liability risk, cross-BU organizational complexity, regulatory scrutiny'],
        ['Revenue Potential', '$37-75M year 1-2 at 5-10% adoption; $750M at full penetration'],
        ['Competitive Moat', 'VERY STRONG -- requires banking core + stablecoin + settlement + POS data'],
        ['Go-to-Market Risk', 'LOW (distribution), MEDIUM (regulatory), HIGH (AI prediction liability)'],
        ['Investor Day Impact', '9/10 -- the $847 dashboard widget is the most demo-able visual'],
        ['Critical Changes', 'Implement Taleb\'s 4 safeguards. Rename to "SmartCash." Clarify yield source. Single P&L owner.'],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ---- IDEA 2: PAY-BY-AGENT ----
doc.add_heading('Idea 2: Pay-by-Agent x402 Commerce', level=2)

doc.add_heading('Expert Analysis', level=3)

expert('Clayton Christensen', 'Disruption Theory',
       'Most strategically important but most speculative. The job -- "pay for something without asking my human" -- is literally impossible on card rails today. x402 is a new-market disruption creating a consumption tier that doesn\'t exist. Risk: $28K/day total global x402 volume. Is this "internet in 1994" or "metaverse in 2022"? The functional job is real, so the former is more apt. The 6M merchant auto-enable is the single most defensible competitive move. In a protocol war, merchant coverage wins. VHS beat Betamax on distribution, not quality.')

expert('Michael Porter', 'Competitive Strategy',
       'WARNING: Stripe is the most dangerous competitor Fiserv has EVER faced here. Stripe controls: MPP, Tempo blockchain, Bridge ($1.1B), Privy, best developer experience, OpenAI/Anthropic relationships, 100+ services at MPP launch. Fiserv controls: x402 Foundation membership (didn\'t create it), 6M merchants, Finxact, FIUSD, INDX. The question: will agentic commerce be won by PROTOCOL ADOPTION (Stripe\'s bet) or MERCHANT COVERAGE (Fiserv\'s bet)? History suggests both matter, but merchant coverage is necessary. Moat is MODERATE -- x402 endpoint is open-source and copyable. Moat is distribution + settlement stack.')

expert('Peter Drucker', 'Management',
       'Challenge: bank-verified agent identity is better for banks and merchants, but WORSE for AI agent developers. Developers want the simplest integration -- which is Stripe\'s one-line API. Fiserv has never been a developer-experience company. Building a competing Agent SDK is a non-trivial organizational challenge. RECOMMENDATION: Partner with an AI lab (Anthropic or OpenAI) to co-build the demo agent. This solves DX and creates a partnership announcement.')

expert('Nassim Taleb', 'Risk Management',
       'Antifragile play: protocol-agnostic support. Do NOT bet solely on x402. Support x402, MPP, AP2, TAP, Agent Pay. If any wins, Fiserv wins. The spending guardrails via Finxact are the strongest risk management feature -- agent payments without guardrails will inevitably produce the "$50K unauthorized purchase" headline. Bank-verified identity with limits makes Fiserv\'s implementation the SAFEST.')

expert('Seth Godin', 'Marketing',
       'If you show an AI agent buying something in 3 seconds -- no passwords, no 3DS, no CVV -- every investor leans forward. This is "iPhone moment" demo quality. But don\'t call it "x402 Commerce." Call it "Fiserv AgentCheckout." The comparison slide "Stripe: one at a time. Fiserv: 6 million overnight" is a mic-drop.')

expert('Donella Meadows', 'Systems Thinking',
       'The 6M auto-enable breaks the chicken-and-egg problem by providing the merchant side unilaterally. Brilliant leverage point. But the feedback loop from enablement to actual agent volume will be SLOW. Do not promise revenue in year 1. Position as infrastructure investment.')

doc.add_heading('Synthesis', level=3)
add_table(
    ['Dimension', 'Assessment'],
    [
        ['Strengths', '6M auto-enable is strongest distribution play, protocol-agnostic gateway is sound, bank-verified identity is differentiated, visionary demo'],
        ['Weaknesses', 'Market barely exists ($28K/day), Stripe is formidable, Agent SDK quality is organizational risk, revenue 3-5 years out'],
        ['Revenue Potential', 'Near-zero year 1. $50-150M/year by 2030 if agentic commerce materializes'],
        ['Competitive Moat', 'MODERATE -- distribution is strong but x402 is open-source. Moat is settlement stack + distribution'],
        ['Go-to-Market Risk', 'HIGH -- depends on market materializing, developer adoption, protocol standardization'],
        ['Investor Day Impact', '10/10 -- live AI agent demo is the only one that creates a "holy sh*t" moment'],
        ['Critical Changes', 'Rename to "AgentCheckout." Make protocol-agnostic. Partner with AI lab. Don\'t promise near-term revenue.'],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ---- IDEA 3: SUPPLIER PAY ----
doc.add_heading('Idea 3: Instant Supplier Pay', level=2)

doc.add_heading('Expert Analysis', level=3)

expert('Clayton Christensen', 'Disruption Theory',
       'Most practical idea. Job-to-be-done is visceral: "Help me pay my suppliers without losing money." Every restaurant owner understands this. BUT: by moving B2B supply payments off card rails onto FIUSD, Fiserv is cannibalizing its own interchange revenue. Every $50K/month that moves from card to FIUSD = $1,000-1,500/month in lost interchange. The FIUSD fee + discount share must be modeled explicitly to ensure unit economics are positive.')

expert('Michael Porter', 'Competitive Strategy',
       'Competitive dynamics are more favorable because this is a VERTICAL solution (restaurant supply chain) not horizontal. Vertical solutions are harder for Stripe/Adyen to replicate because they need domain-specific data. Clover has this data -- that\'s the moat. Revenue: 500K restaurants x $50K/month supplies x 30bps = $90M/year at full penetration. Plus 10% of 2% discount on $50K x 500K = $600M/year. At 5% year-1 adoption = ~$35M.')

expert('Peter Drucker', 'Management',
       'Requires a NEW relationship Fiserv doesn\'t have: relationships with food distributors. Will Sysco and US Foods (50% of US restaurant distribution) accept FIUSD? Without major distributors, this works only with small local suppliers. RECOMMENDATION: Call Sysco and US Foods BEFORE building. If yes, this jumps to priority 1. If no, addressable market shrinks dramatically.')

expert('Nassim Taleb', 'Risk Management',
       'Best risk profile of all four. Downside is bounded: bad prediction = slightly suboptimal order (too much/little of an ingredient). Compare to Idea 1 (can\'t make payroll) or Idea 2 ($50K unauthorized spend). The early-payment discount capture is genuinely asymmetric: 2% for paying 20 days early = ~36% annualized return. Free money restaurants leave on the table.')

expert('Seth Godin', 'Marketing',
       'Compelling for a DIFFERENT audience. For merchants: "Your AI saved $1,200 this month." For investors, elevate: "We are replacing the $2.4 trillion B2B card payment market with zero-fee stablecoin rails." Best as SUPPORTING demo, not lead.')

doc.add_heading('Synthesis', level=3)
add_table(
    ['Dimension', 'Assessment'],
    [
        ['Strengths', 'Quantifiable pain point, bounded downside, 36% annualized discount capture, vertical data moat, favorable risk profile'],
        ['Weaknesses', 'Requires distributor buy-in, cannibalizes interchange, prediction accuracy on inferred data, restaurant-only initially'],
        ['Revenue Potential', '$35M year 1 at 5% adoption; $90M txn fees + $600M discount share at full penetration'],
        ['Competitive Moat', 'STRONG for restaurant vertical (Clover data). WEAK for generalization'],
        ['Go-to-Market Risk', 'MEDIUM -- depends on distributor willingness. Cannibalization math must work'],
        ['Investor Day Impact', '7/10 -- compelling feature demo, best as supporting evidence of platform versatility'],
        ['Critical Changes', 'Validate distributor willingness. Model interchange cannibalization. Position as "first vertical app of stablecoin infrastructure."'],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ---- IDEA 4: CROSS-BORDER ----
doc.add_heading('Idea 4: Cross-Border Instant Settlement', level=2)

doc.add_heading('Expert Analysis', level=3)

expert('Clayton Christensen', 'Disruption Theory',
       'Largest market but most competitive. NOT a blue ocean -- red ocean with Stripe (Bridge), Worldpay (BVNK), Wise, Revolut fighting over cross-border. Fiserv\'s unique advantage: EMBEDDED cross-border settlement within CommerceHub. Merchant doesn\'t go to Wise. The transaction is detected and settled optimally within existing workflow. This is a convenience disruption, not cost disruption -- and convenience often wins because it requires zero behavior change.')

expert('Michael Porter', 'Competitive Strategy',
       'The $190T market is misleading. Addressable = cross-border e-commerce ($6.3T, growing 25%/year) + SMB B2B (~$2-5T). Competitive intensity: HIGH. Supplier power: HIGH (FX liquidity providers). Substitutes: ABUNDANT. The unique positioning: Fiserv is the ONLY company offering cross-border settlement as a native feature of an existing orchestration platform. Revenue: $10-50B CommerceHub cross-border volume x 70bps = $70-350M/year.')

expert('Nassim Taleb', 'Risk Management',
       'Best risk/reward ratio of all four. FX conversion and cross-border settlement are WELL-UNDERSTOOD operations. The stablecoin rail makes it faster and cheaper but the operation is not novel. Implementation risk: lower. Regulatory framework: clearer. Failure modes: well-mapped. Primary risk: FX rate slippage during the 30-second lock window. At $10B+ volume, even 5bps average slippage = $5M/year loss.')

expert('Kim & Mauborgne', 'Blue Ocean Strategy',
       'Fiserv is NOT creating blue ocean here -- offering a BETTER solution in existing red ocean. The blue ocean element: "cross-border payments the merchant doesn\'t know are cross-border." If CommerceHub automatically detects, converts, and settles with zero merchant intervention -- merchant just sees "payment received, USD settled" -- THAT is new. The invisibility of cross-border complexity is the blue ocean.')

doc.add_heading('Synthesis', level=3)
add_table(
    ['Dimension', 'Assessment'],
    [
        ['Strengths', 'Largest addressable market ($6.3T e-comm), well-understood operation, best risk/reward, 90% cost reduction is dramatic, highly demo-able'],
        ['Weaknesses', 'Red ocean (Wise, Bridge, BVNK), regulatory complexity for stablecoin cross-border, FX slippage risk, not category-creating'],
        ['Revenue Potential', '$70-350M/year at 70bps on $10-50B cross-border CommerceHub volume. Most certain revenue model'],
        ['Competitive Moat', 'MODERATE -- embedded in CommerceHub gives distribution advantage, but tech is commoditized'],
        ['Go-to-Market Risk', 'MEDIUM -- regulatory approval is gating factor. Technology risk is low'],
        ['Investor Day Impact', '8/10 -- side-by-side comparison ($47.50 vs $5.00) is immediately compelling'],
        ['Critical Changes', 'Make invisible (auto-route, no merchant opt-in). Validate regulatory posture. Use $6.3T not $190T for TAM.'],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ---- OVERALL RANKING ----
doc.add_heading('Overall Business Panel Ranking', level=2)

doc.add_heading('Expert Rankings', level=3)
add_table(
    ['Expert', 'Framework', '#1', '#2', '#3', '#4'],
    [
        ['Christensen', 'Disruption potential', 'Pay-by-Agent', 'Yield Sweep', 'Supplier Pay', 'Cross-Border'],
        ['Porter', 'Competitive defensibility', 'Yield Sweep', 'Supplier Pay', 'Cross-Border', 'Pay-by-Agent'],
        ['Drucker', 'Organizational feasibility', 'Yield Sweep', 'Cross-Border', 'Supplier Pay', 'Pay-by-Agent'],
        ['Taleb', 'Risk-adjusted return', 'Cross-Border', 'Yield Sweep', 'Supplier Pay', 'Pay-by-Agent'],
        ['Godin', 'Investor Day impact', 'Pay-by-Agent', 'Yield Sweep', 'Cross-Border', 'Supplier Pay'],
        ['Collins', 'Flywheel strength', 'Yield Sweep', 'Pay-by-Agent', 'Cross-Border', 'Supplier Pay'],
        ['Meadows', 'System leverage', 'Pay-by-Agent', 'Cross-Border', 'Yield Sweep', 'Supplier Pay'],
    ],
    col_widths=[1.3, 1.8, 1.5, 1.5, 1.5, 1.5]
)
doc.add_paragraph()

doc.add_heading('Consensus Ranking', level=3)
add_table(
    ['Rank', 'Idea', 'Score', 'Investor Day Role'],
    [
        ['1', 'Merchant Yield Sweep', '9/10', 'LEAD DEMO -- immediate, tangible merchant value'],
        ['2', 'Pay-by-Agent x402', '10/10', 'VISIONARY DEMO -- future of commerce'],
        ['3', 'Cross-Border Settlement', '8/10', 'SCALE DEMO -- massive TAM, clear cost savings'],
        ['4', 'Instant Supplier Pay', '7/10', 'SUPPORTING DEMO -- platform versatility'],
    ],
    col_widths=[0.5, 2.5, 0.7, 3.5]
)
doc.add_paragraph()

p('All experts agree: Build all 4. Together they tell a story no single idea tells alone. Yield Sweep proves immediate value. Pay-by-Agent proves visionary positioning. Cross-Border proves massive TAM. Supplier Pay proves B2B versatility. Remove any one and the Investor Day narrative weakens.', bold=True, color=ORANGE)

doc.add_heading('Panel Consensus Recommendations', level=3)
for item in [
    'Yield Sweep: Clarify yield source. Implement 4 safeguards (hard floor, instant override, Fiserv backstop, gradual ramp). Rename to "SmartCash." Appoint single cross-BU P&L owner.',
    'Pay-by-Agent: Make protocol-agnostic (x402 + MPP + AP2). Rename to "Fiserv AgentCheckout." Explore AI lab partnership. Don\'t promise near-term revenue -- position as infrastructure bet.',
    'Supplier Pay: Validate distributor willingness (call Sysco/US Foods). Model interchange cannibalization explicitly. Position as "first vertical application of stablecoin infrastructure."',
    'Cross-Border: Validate regulatory posture with legal. Make stablecoin routing invisible (automatic, not opt-in). Use $6.3T e-commerce TAM, not $190T total.',
    'Investor Day Flow: Lead with Yield Sweep (immediate value) -> Pay-by-Agent (visionary) -> Supplier Pay + Cross-Border combined (platform scale). Arc: "value today, future of commerce, massive TAM."',
]:
    bullet(item)

doc.add_page_break()

# ============================================================
# SECTION 2: SPEC PANEL
# ============================================================
doc.add_heading('Section 2: Technical Specification Panel Review', level=1)
p('Mode: CRITIQUE  |  Focus: Architecture + Testing + Reliability', bold=True, color=ORANGE)
p('Experts: Martin Fowler (Architecture), Michael Nygard (Production Reliability), Sam Newman (Distributed Systems), Karl Wiegers (Requirements), Gojko Adzic (Specification by Example), Lisa Crispin (Testing), Kelsey Hightower (Cloud/Operations), Gregor Hohpe (Integration Patterns)')

doc.add_page_break()

# ---- PROTO 1 SPEC ----
doc.add_heading('Prototype 1: Merchant Yield Sweep -- Technical Review', level=2)

doc.add_heading('Architecture Assessment', level=3)

expert('Martin Fowler', 'Architecture',
       'Architecture is clean pipeline: data ingestion -> prediction -> action -> settlement. Three concerns: (1) Tight coupling between prediction and execution. Need a DECISION GATE between predictor and Finxact executor -- a separate service that validates safeguards before money moves. (2) Cron-based scheduling is fragile for demo -- add a /sweeps/evaluate endpoint for on-demand triggering. (3) PostgreSQL schema must be designed in week 1. Recommended 6 entities: merchants, transactions, predictions, sweep_decisions, sweep_executions, yield_accruals.')

expert('Michael Nygard', 'Production Reliability',
       '#1 demo risk: Finxact API call fails during live demo. Sandboxes have lower SLAs. CRITICAL: Build a "demo mode" toggle. When enabled, uses pre-computed predictions and pre-approved sweep decisions for the demo merchant. Real API calls still execute underneath, but if anything fails, seamlessly falls back to pre-computed path. Audience never sees the difference. Also pre-validate all demo accounts before going on stage.')

expert('Sam Newman', 'Distributed Systems',
       'Don\'t use real CommerceHub webhooks for the prototype. Build a simulated settlement event generator that produces realistic transaction data. Removes dependency on CommerceHub sandbox availability and gives full control over demo data.')

expert('Lisa Crispin', 'Testing',
       'Testing priorities: Week 3-4: Integration tests for Finxact API calls. Week 5: End-to-end demo rehearsal (run 10 times, fix every failure). Week 6: Chaos testing -- kill Finxact connection mid-sweep, kill PostgreSQL, verify demo mode fallback works.')

doc.add_heading('Feasibility & Score', level=3)
add_table(
    ['Component', 'Complexity', 'Risk', 'Weeks'],
    [
        ['Settlement simulator', 'Low', 'Low', '0.5'],
        ['Prediction model', 'Medium', 'Medium', '1.5'],
        ['Decision gate + safeguards', 'Medium', 'Low', '1'],
        ['Finxact integration', 'High', 'HIGH', '1.5'],
        ['INDX simulation', 'Low', 'Low', '0.5'],
        ['React dashboard', 'Medium', 'Low', '1.5'],
        ['Demo mode + rehearsal', 'Low', 'Medium', '0.5'],
        ['TOTAL', '', '', '7 weeks'],
    ]
)
doc.add_paragraph()
p('Verdict: FEASIBLE with 2-3 engineers, but tight. Finxact integration is critical path.', bold=True, color=GREEN)

doc.add_heading('Missing / Underspecified', level=3)
for item in [
    'Yield source: Where does FIUSD yield come from? Simulate 4.2% APY for prototype. Prepare answer for Investor Day Q&A.',
    'Merchant onboarding flow: "One toggle" mentioned but UX not defined (terms, risk tolerance, bank verification).',
    'Tax implications: Yield may be taxable income. Irrelevant for prototype, critical for production.',
    'Multi-tenant isolation: Row-level security or strict FK constraints in PostgreSQL.',
]:
    bullet(item)

doc.add_page_break()

# ---- PROTO 2 SPEC ----
doc.add_heading('Prototype 2: Pay-by-Agent x402 -- Technical Review', level=2)

doc.add_heading('Architecture Assessment', level=3)

expert('Martin Fowler', 'Architecture',
       'Most architecturally complex prototype. Three concerns: (1) Separate Gateway from Verifier from Settler -- different failure characteristics need different handling. Gateway handles HTTP, Settler handles on-chain (slow/unreliable). (2) EIP-3009 vs Permit2: pick ONE for prototype. EIP-3009 on USDC/FIUSD is sufficient. (3) Claude demo agent is highest-risk component. LLMs are non-deterministic. If agent hallucinates during live demo, the demo fails publicly. Must be heavily constrained.')

expert('Michael Nygard', 'Production Reliability',
       'SIX potential failure points in a single transaction: LLM hallucination, gateway down, wallet key error, crypto library bug, Solana devnet down, INDX mock failure. CRITICAL: Solana devnet has no SLA and goes down periodically. RECOMMENDATION: Run LOCAL Solana validator (solana-test-validator) on the demo machine. Pre-deploy tokens, pre-fund wallets. Zero external dependencies. Same for Base: use Hardhat local node.')

expert('Sam Newman', 'Distributed Systems',
       'Define the Agent SDK TypeScript interface in week 1. It drives everything: fetchWithPayment(url, options) -> PaymentResponse with status, receipt, resource, settlementTxId, amount, token, chain. Define the contract first, implement second.')

expert('Gojko Adzic', 'Specification by Example',
       'Demo agent MUST be scripted, not open-ended. Hardcode the product catalog (10-20 items). Constrain tool_use to 3 tools: search_products, get_product, purchase. The only "real" part should be the x402 payment flow. Define exact demo scenario as an automated test.')

doc.add_heading('Feasibility & Score', level=3)
add_table(
    ['Component', 'Complexity', 'Risk', 'Weeks'],
    [
        ['x402 Gateway (HTTP 402 spec)', 'High', 'Medium', '2'],
        ['Payment Verifier (EIP-3009)', 'High', 'HIGH', '1.5'],
        ['Local Solana validator setup', 'Medium', 'Medium', '0.5'],
        ['Agent SDK (TypeScript only)', 'Medium', 'Low', '1.5'],
        ['Claude demo agent', 'Medium', 'HIGH', '1'],
        ['INDX simulation', 'Low', 'Low', '0.5'],
        ['React dashboard', 'Medium', 'Low', '1'],
        ['Demo rehearsal + hardening', 'Low', 'Medium', '1'],
        ['TOTAL', '', '', '9.5 weeks'],
    ]
)
doc.add_paragraph()
p('Verdict: TIGHT for 7 weeks. Requires scope cuts: Cut Python SDK (TypeScript only). Cut Permit2 (EIP-3009 only). Cut Base chain (Solana only). Allocate 4 engineers, not 3.', bold=True, color=RED)

doc.add_heading('Critical Scope Cuts for 7-Week Feasibility', level=3)
for item in [
    'CUT Python SDK -- TypeScript only for Investor Day. Add Python post-demo.',
    'CUT Permit2 -- EIP-3009 only (USDC/FIUSD).',
    'CUT Base chain -- Solana only for demo. Add Base post-demo.',
    'USE local Solana validator -- not devnet. Eliminates #1 demo reliability risk.',
    'CONSTRAIN Claude agent -- 3 tools only, hardcoded product catalog, scripted scenario.',
    'ALLOCATE 4 engineers -- this prototype needs the extra capacity.',
]:
    bullet(item)

doc.add_page_break()

# ---- PROTO 3 SPEC ----
doc.add_heading('Prototype 3: Instant Supplier Pay -- Technical Review', level=2)

doc.add_heading('Architecture Assessment', level=3)

expert('Martin Fowler', 'Architecture',
       'Most straightforward architecture. Clean pipeline: sales -> prediction -> PO -> payment -> settlement. Key hidden complexity: BILL OF MATERIALS (BOM) mapping. Clover tracks what was SOLD (menu items), not what was CONSUMED (ingredients). Need a BOM mapping each menu item to ingredient components. For demo, hardcode for "Mario\'s Pizzeria" with 15-20 menu items and 20 ingredients. Also: cut approval workflow -- auto-approve everything for demo, show config UI only.')

expert('Michael Nygard', 'Production Reliability',
       'Demo reliability risk: LOWEST of all four. Entire flow can be simulated without external dependencies. Only external: Finxact API for B2B transfer (apply demo mode fallback). Key demo visual: the "3 seconds" settlement claim needs a LIVE timer/animation. Build a progress bar showing each step completing in real time.')

expert('Sam Newman', 'Distributed Systems',
       'Critical question: Does Finxact API support atomic cross-account transfers? Or two separate calls (debit merchant, credit supplier)? If two calls, need compensation mechanism -- if debit succeeds but credit fails, must reverse the debit. Wrap in try/except with rollback.')

doc.add_heading('Feasibility & Score', level=3)
add_table(
    ['Component', 'Complexity', 'Risk', 'Weeks'],
    [
        ['Sales simulator + BOM', 'Medium', 'Low', '1'],
        ['Ingredient depletion model', 'Medium', 'Medium', '1.5'],
        ['Supplier catalog + PO engine', 'Low', 'Low', '1'],
        ['Finxact B2B payment', 'Medium', 'Medium', '1'],
        ['INDX simulation', 'Low', 'Low', '0.5'],
        ['React dashboard', 'Medium', 'Low', '1.5'],
        ['Demo scenario + rehearsal', 'Low', 'Low', '0.5'],
        ['TOTAL', '', '', '7 weeks'],
    ]
)
doc.add_paragraph()
p('Verdict: FEASIBLE with 2-3 engineers. Most comfortable timeline of all four prototypes.', bold=True, color=GREEN)

doc.add_heading('Missing / Key Additions', level=3)
for item in [
    'Bill of Materials (BOM) mapping -- the critical data model connecting POS sales to ingredient consumption. Hardcode as JSON config.',
    'Supplier API mock -- simple Express server accepting POs and returning confirmations.',
    'Settlement timer animation -- the "2.7 seconds" visual is the money shot for this demo.',
    'Pre-populate 30 days of simulated sales data at startup so prediction model has history.',
]:
    bullet(item)

doc.add_page_break()

# ---- PROTO 4 SPEC ----
doc.add_heading('Prototype 4: Cross-Border Instant Settlement -- Technical Review', level=2)

doc.add_heading('Architecture Assessment', level=3)

expert('Martin Fowler', 'Architecture',
       'Architecture is sound but FX layer has hidden complexity. (1) The 30-second rate lock means Fiserv takes FX risk. For prototype: use FIXED simulated rate. (2) Cross-border detection is underspecified -- for prototype, hardcode a "demo buyer" with explicit currency. Don\'t build IP geolocation or BIN lookup. (3) Reconciliation engine is overkill for demo -- build the VISUAL showing 3 steps completing in sequence, skip the ledger matching logic.')

expert('Michael Nygard', 'Production Reliability',
       'Same Solana devnet risk as Prototype 2 -- use local validator. Demo flow timing challenge: need side-by-side comparison LIVE. Show card route (static: "$47.50, 3 days") then stablecoin route executing in real time with step-by-step animation. Pre-compute expected results so animation runs smoothly even if API is slow.')

expert('Gregor Hohpe', 'Integration Patterns',
       'Classic Content-Based Router pattern. THREE corridors spec\'d but demo only needs ONE. Build MXN->USD for demo. Show EUR and GBP as config options in UI but don\'t implement conversion logic for those corridors. Reduces testing surface by 66%.')

expert('Kelsey Hightower', 'Cloud/Operations',
       'Cross-border has regulatory implications. Add a PLACEHOLDER: log line "OFAC screening: PASSED (simulated)" in transaction detail view. Shows investors Fiserv is thinking about compliance without implementing it.')

doc.add_heading('Feasibility & Score', level=3)
add_table(
    ['Component', 'Complexity', 'Risk', 'Weeks'],
    [
        ['Demo buyer simulation', 'Low', 'Low', '0.5'],
        ['FX conversion (MXN only)', 'Medium', 'Low', '1'],
        ['Rate lock (simplified)', 'Low', 'Low', '0.5'],
        ['Local Solana FIUSD transfer', 'Medium', 'Medium', '1'],
        ['INDX simulation', 'Low', 'Low', '0.5'],
        ['Side-by-side dashboard', 'Medium', 'Low', '1.5'],
        ['Settlement animation', 'Medium', 'Medium', '1'],
        ['Demo rehearsal', 'Low', 'Low', '0.5'],
        ['TOTAL', '', '', '6.5 weeks'],
    ]
)
doc.add_paragraph()
p('Verdict: FEASIBLE with 2-3 engineers. Second most comfortable timeline.', bold=True, color=GREEN)

doc.add_heading('Key Simplifications for Demo', level=3)
for item in [
    'Single corridor only (MXN->USD) -- show EUR/GBP as config options.',
    'Fixed FX rate -- simulate rate lock visually but don\'t handle real volatility.',
    'Local Solana validator -- not devnet.',
    'Hardcode demo buyer ("Carlos Rodriguez", MXN, 17,500 pesos).',
    'Skip reconciliation engine -- build the visual, not the ledger matching.',
    'Add OFAC compliance stub -- "Screening: PASSED (simulated)" in transaction detail.',
]:
    bullet(item)

doc.add_page_break()

# ---- CROSS-CUTTING ----
doc.add_heading('Cross-Cutting Technical Recommendations', level=2)

doc.add_heading('Shared Infrastructure (Build First)', level=3)
p('These four prototypes share significant infrastructure. Extract shared components to save 2-3 weeks:', bold=True)

add_table(
    ['Component', 'Description', 'Used By'],
    [
        ['Finxact Client Library', 'Python SDK wrapping all Finxact REST API calls (accounts, positions, transfers)', 'P1, P3, P4'],
        ['INDX Simulator', 'Mock service: POST /settle -> USD settlement confirmation with 1-3s latency', 'All 4'],
        ['Local Solana Setup', 'Script: start validator, deploy FIUSD/USDC tokens, pre-fund wallets', 'P2, P4'],
        ['React Component Library', 'Earnings widgets, settlement animations, transaction tables', 'All 4'],
        ['Demo Mode Pattern', 'Universal fallback to pre-computed results if any external dependency fails', 'All 4'],
    ],
    col_widths=[1.8, 3.5, 1.0]
)
doc.add_paragraph()

p('Monorepo structure:', bold=True)
p('/shared/\n  /finxact-client/\n  /indx-simulator/\n  /solana-setup/\n  /react-components/\n  /demo-mode/\n/prototype-1-yield-sweep/\n/prototype-2-agent-pay/\n/prototype-3-supplier-pay/\n/prototype-4-cross-border/', size=10)

doc.add_heading('Demo Reliability Pattern (All Prototypes)', level=3)
p('Every prototype must implement the same pattern: attempt real API call with 5-second timeout. If it fails, seamlessly fall back to pre-computed result. Log the fallback but don\'t show it to the audience. Register fallbacks for every external dependency at startup.')

doc.add_heading('Investor Day Technical Setup', level=3)
for item in [
    'Single demo machine running all 4 prototypes locally. No cloud dependencies.',
    'Local Solana validator pre-started with all tokens deployed and wallets funded.',
    'All databases pre-seeded with 30+ days of simulated data.',
    'All demo mode fallbacks registered and tested.',
    'Network independence: entire demo must work with WiFi OFF. Conference WiFi is unreliable.',
    'Backup machine: identical setup on a second laptop. If primary dies, swap in 60 seconds.',
    'Screen recording: record demo in advance as backup video. If both machines fail, play the video.',
    'Pre-demo health check script: curl all 4 /health endpoints, check Solana validator, PostgreSQL, Redis.',
]:
    bullet(item)

doc.add_page_break()

# ---- SPEC PANEL OVERALL SCORES ----
doc.add_heading('Spec Panel Overall Scores', level=2)
add_table(
    ['Dimension', 'P1: Yield Sweep', 'P2: Agent Pay', 'P3: Supplier Pay', 'P4: Cross-Border'],
    [
        ['Architecture soundness', '8/10', '7/10', '9/10', '8/10'],
        ['API design quality', '7/10', '8/10', '7/10', '7/10'],
        ['Security posture', '7/10', '6/10', '8/10', '7/10'],
        ['Demo reliability', '7/10', '5/10 (highest risk)', '9/10 (lowest risk)', '7/10'],
        ['Feasibility in 7 weeks', '8/10', '6/10 (tightest)', '9/10 (most comfortable)', '8/10'],
        ['Spec completeness', '7/10', '6/10', '7/10', '7/10'],
        ['OVERALL', '7.3/10', '6.3/10', '8.2/10', '7.3/10'],
    ],
    col_widths=[1.8, 1.5, 1.8, 1.8, 1.5]
)

doc.add_page_break()

# ============================================================
# SECTION 3: CONSOLIDATED ACTION ITEMS
# ============================================================
doc.add_heading('Section 3: Consolidated Action Items', level=1)

doc.add_heading('Week 1 Priority Actions (All Teams)', level=2)
for item in [
    'Build shared infrastructure first: Finxact client, INDX simulator, Solana setup, React components, demo mode pattern.',
    'Define ALL API interfaces and data models before writing implementation code.',
    'Set up local Solana validator with FIUSD/USDC token contracts deployed.',
    'Design PostgreSQL schemas for all four prototypes.',
    'Set up monorepo structure: /shared/ + /prototype-N/ directories.',
    'Validate Finxact sandbox API access and document available endpoints.',
]:
    bullet(item)

doc.add_heading('Team Allocation (Updated per Spec Panel)', level=2)
add_table(
    ['Prototype', 'Engineers', 'Rationale'],
    [
        ['1. Yield Sweep', '2-3', 'Feasible. Finxact integration is critical path.'],
        ['2. Agent Pay', '4 (increased from 3)', 'Tightest timeline, highest risk. Extra capacity needed for x402 + crypto + agent.'],
        ['3. Supplier Pay', '2-3', 'Most comfortable timeline. Standard web app + Finxact integration.'],
        ['4. Cross-Border', '2-3', 'Comfortable with scope simplifications (single corridor, fixed FX, local Solana).'],
        ['Shared Infrastructure', '1 (first 2 weeks)', 'Dedicated engineer builds shared components before prototype work begins.'],
    ],
    col_widths=[1.5, 1.2, 4.5]
)
doc.add_paragraph()
p('Total: 12-14 engineers (1 shared infra + 11-13 across 4 prototypes)', bold=True)

doc.add_heading('Scope Cuts for Demo Feasibility', level=2)
add_table(
    ['Prototype', 'Cut from Demo', 'Reason', 'Add Post-Investor Day'],
    [
        ['2. Agent Pay', 'Python SDK', 'Saves 0.5 weeks, reduces testing', 'Yes -- add in sprint 2'],
        ['2. Agent Pay', 'Permit2 support', 'EIP-3009 sufficient for USDC/FIUSD', 'Yes'],
        ['2. Agent Pay', 'Base chain', 'Solana only for demo. Local validator.', 'Yes'],
        ['3. Supplier Pay', 'Approval workflows', 'Auto-approve for demo, show config UI only', 'Yes'],
        ['4. Cross-Border', 'EUR/GBP corridors', 'MXN only for demo. Show others as config.', 'Yes'],
        ['4. Cross-Border', 'Reconciliation engine', 'Build visual animation, skip ledger matching', 'Yes'],
        ['4. Cross-Border', 'Geolocation/BIN detection', 'Hardcode demo buyer', 'Yes'],
    ],
    col_widths=[1.3, 1.8, 2.5, 1.5]
)
doc.add_paragraph()

doc.add_heading('Key Additions from Panel Reviews', level=2)
add_table(
    ['Addition', 'Prototype', 'Why'],
    [
        ['Decision Gate service', 'P1', 'Validates safeguards independently before money moves. Auditable.'],
        ['Demo mode toggle', 'All 4', 'Pre-computed fallback if any external dependency fails during live demo.'],
        ['Settlement simulator', 'P1', 'Replaces CommerceHub webhook dependency with controlled data.'],
        ['Local Solana validator', 'P2, P4', 'Eliminates Solana devnet SLA risk. Fully offline demo.'],
        ['Bill of Materials mapping', 'P3', 'Connects POS sales to ingredient consumption. Hardcode as JSON.'],
        ['Settlement timer animation', 'P3, P4', 'The "3 seconds" claim needs live visual proof.'],
        ['OFAC compliance stub', 'P4', 'Shows regulatory awareness without implementation. "Screening: PASSED (simulated)."'],
        ['Backup demo video', 'All 4', 'Pre-recorded fallback if both demo machines fail.'],
    ],
    col_widths=[2.0, 0.7, 4.5]
)

# ============================================================
# HEADERS / FOOTERS / SAVE
# ============================================================
for section in doc.sections:
    header = section.header
    hp = header.paragraphs[0]
    hp.text = 'Fiserv Digital Pay | Expert Panel Evaluations | CONFIDENTIAL'
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in hp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = 'May 2026 | Business Panel + Spec Panel Analysis'
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'

output = '/Users/ajnarasi/Documents/Work/Projects/stablecoin/deliverables/Fiserv_Panel_Evaluations.docx'
doc.save(output)
print(f'Document saved to: {output}')
